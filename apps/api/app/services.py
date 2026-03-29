from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi import UploadFile

from .document_processor import DocumentProcessor
from .executor import build_execution_adapter
from .models import (
    ArtifactExportPayload,
    ArtifactExportResult,
    ArtifactRecord,
    ArtifactReviewPayload,
    CaseCreate,
    CaseDetail,
    CaseRecord,
    DocumentDetail,
    DocumentRecord,
    ReviewRecord,
    TaskCreate,
    TaskRecord,
    TimelineItem,
    WorkflowRunRecord,
    utc_now,
)
from .store import JsonStore
from .workflow_engine import WorkflowEngine


class WorkspaceService:
    def __init__(self, store: JsonStore, executor_mode: str = "local", executor_command: str | None = None):
        self.store = store
        self.adapter = build_execution_adapter(store, executor_mode, command=executor_command)
        self.document_processor = DocumentProcessor()

    def create_case(self, payload: CaseCreate) -> CaseRecord:
        workspace = self.store.ensure_case_workspace(payload.case_code)
        case = CaseRecord(id=self.store.next_id("case"), **payload.model_dump())
        self.store.create("cases", case.model_dump())
        self._append_timeline(case.id, "case_created", f"Case workspace initialized: {workspace.name}")
        return case

    def list_cases(self) -> list[CaseRecord]:
        return [CaseRecord(**item) for item in self.store.list("cases")]

    def get_case(self, case_id: str) -> CaseDetail | None:
        case = self.store.get("cases", case_id)
        if not case:
            return None
        tasks = [
            self._refresh_task(TaskRecord(**item))
            for item in self.store.filter_by("tasks", case_id=case_id)
        ]
        return CaseDetail(
            **case,
            documents=[DocumentRecord(**item) for item in self.store.filter_by("documents", case_id=case_id)],
            tasks=tasks,
            artifacts=[ArtifactRecord(**item) for item in self.store.filter_by("artifacts", case_id=case_id)],
            workflow_runs=[WorkflowRunRecord(**item) for item in self.store.filter_by("workflow_runs", case_id=case_id)],
            review_records=[ReviewRecord(**item) for item in self.store.filter_by("review_records", case_id=case_id)],
            timeline=[TimelineItem(**item) for item in self.store.filter_by("timelines", case_id=case_id)],
        )

    async def add_document(self, case_id: str, category: str, file: UploadFile) -> DocumentRecord:
        case = self.store.get("cases", case_id)
        if not case:
            raise KeyError(case_id)

        target_dir = self.store.ensure_case_subdir(case["case_code"], category)
        file_path = target_dir / file.filename
        content = await file.read()
        file_path.write_bytes(content)

        extracted_path = file_path.with_suffix(".md")
        text = self.document_processor.extract(file.filename, content, target_dir)
        extracted_path.write_text(text, encoding="utf-8")

        document = DocumentRecord(
            id=self.store.next_id("doc"),
            case_id=case_id,
            file_name=file.filename,
            file_type=file.content_type or "application/octet-stream",
            storage_path=str(file_path),
            extracted_text_path=str(extracted_path),
            category=category,
            processing_status="processed",
        )
        self.store.create("documents", document.model_dump())
        self._append_timeline(case_id, "document_uploaded", f"Document uploaded: {file.filename}")
        return document

    def get_document(self, document_id: str) -> DocumentDetail | None:
        record = self.store.get("documents", document_id)
        if not record:
            return None
        document = DocumentRecord(**record)
        extracted_text = Path(document.extracted_text_path).read_text(encoding="utf-8")
        return DocumentDetail(**document.model_dump(), extracted_text=extracted_text)

    def create_task(self, case_id: str, payload: TaskCreate) -> TaskRecord:
        case = self.store.get("cases", case_id)
        if not case:
            raise KeyError(case_id)

        documents = self._get_documents(case_id, payload.document_ids)
        workflow_run = WorkflowRunRecord(
            id=self.store.next_id("run"),
            case_id=case_id,
            workflow_type=payload.task_type,
            title=payload.title,
            status="running",
        )
        task = TaskRecord(
            id=self.store.next_id("task"),
            case_id=case_id,
            workflow_run_id=workflow_run.id,
            status="pending",
            **payload.model_dump(),
        )
        workflow_run = workflow_run.model_copy(update={"task_ids": [task.id]})
        self.store.create("workflow_runs", workflow_run.model_dump())

        executed_task, artifact = self.adapter.submit(case, task, documents)
        self.store.create("tasks", executed_task.model_dump())
        self._sync_workflow_run_from_task(executed_task)

        if artifact:
            self._store_artifact_once(artifact)
            self._append_timeline(case_id, "task_completed", f"Task completed: {payload.title}")
        elif executed_task.status == "failed":
            self._append_timeline(case_id, "task_failed", f"Task failed: {payload.title}")
        else:
            self._append_timeline(case_id, "task_submitted", f"Task submitted: {payload.title}")
        return executed_task

    def retry_task(self, task_id: str) -> TaskRecord:
        existing = self.store.get("tasks", task_id)
        if not existing:
            raise KeyError(task_id)
        case = self.store.get("cases", existing["case_id"])
        if not case:
            raise KeyError(existing["case_id"])

        current = TaskRecord(**existing)
        documents = self._get_documents(current.case_id, current.document_ids)
        retry_task = current.model_copy(
            update={
                "status": "pending",
                "attempts": current.attempts + 1,
                "artifact_ids": [],
                "error_message": "",
                "logs": current.logs + [f"{utc_now()} retry requested"],
                "updated_at": utc_now(),
            }
        )
        executed_task, artifact = self.adapter.submit(case, retry_task, documents)
        self.store.update("tasks", task_id, executed_task.model_dump())
        self._sync_workflow_run_from_task(executed_task)

        if artifact:
            self._store_artifact_once(artifact)
            self._append_timeline(current.case_id, "task_retried", f"Task retried successfully: {current.title}")
        elif executed_task.status == "running":
            self._append_timeline(current.case_id, "task_retried", f"Task resubmitted: {current.title}")
        else:
            self._append_timeline(current.case_id, "task_retried", f"Task retry failed: {current.title}")
        return executed_task

    def get_task(self, task_id: str) -> TaskRecord | None:
        record = self.store.get("tasks", task_id)
        if not record:
            return None
        task = TaskRecord(**record)
        return self._refresh_task(task)

    def cancel_task(self, task_id: str) -> TaskRecord:
        record = self.store.get("tasks", task_id)
        if not record:
            raise KeyError(task_id)
        task = TaskRecord(**record)
        canceled = self.adapter.cancel(task)
        self.store.update("tasks", task_id, canceled.model_dump())
        self._sync_workflow_run_from_task(canceled)
        self._append_timeline(canceled.case_id, "task_canceled", f"Task canceled: {canceled.title}")
        return canceled

    def get_artifact(self, artifact_id: str) -> ArtifactRecord | None:
        record = self.store.get("artifacts", artifact_id)
        return ArtifactRecord(**record) if record else None

    def review_artifact(self, artifact_id: str, payload: ArtifactReviewPayload) -> ArtifactRecord:
        record = self.store.get("artifacts", artifact_id)
        if not record:
            raise KeyError(artifact_id)

        artifact = ArtifactRecord(**record).model_copy(
            update={
                "review_status": payload.action,
                "review_comment": payload.comment,
                "reviewed_by": payload.reviewer_name,
                "reviewed_at": utc_now(),
            }
        )
        self.store.update("artifacts", artifact_id, artifact.model_dump())

        review_record = ReviewRecord(
            id=self.store.next_id("review"),
            case_id=artifact.case_id,
            artifact_id=artifact.id,
            action=payload.action,
            reviewer_name=payload.reviewer_name,
            comment=payload.comment,
        )
        self.store.create("review_records", review_record.model_dump())
        self._append_timeline(
            artifact.case_id,
            "artifact_reviewed",
            f"{payload.reviewer_name} marked artifact `{artifact.title}` as {payload.action}",
        )
        return artifact

    def list_review_records(self, artifact_id: str) -> list[ReviewRecord]:
        return [ReviewRecord(**item) for item in self.store.filter_by("review_records", artifact_id=artifact_id)]

    def export_artifact(self, artifact_id: str, payload: ArtifactExportPayload) -> ArtifactExportResult:
        artifact = self.get_artifact(artifact_id)
        if not artifact:
            raise KeyError(artifact_id)

        case = self.store.get("cases", artifact.case_id)
        if not case:
            raise KeyError(artifact.case_id)

        export_dir = self.store.ensure_case_subdir(case["case_code"], "10 - Reports")
        if payload.format == "md":
            export_path = export_dir / f"{artifact.title}.export.md"
            export_path.write_text(artifact.content, encoding="utf-8")
        else:
            export_path = export_dir / f"{artifact.title}.docx"
            document = Document()
            for line in artifact.content.splitlines():
                if line.startswith("# "):
                    document.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    document.add_heading(line[3:].strip(), level=2)
                elif line.startswith("- "):
                    document.add_paragraph(line[2:].strip(), style="List Bullet")
                else:
                    document.add_paragraph(line)
            document.save(export_path)

        self._append_timeline(artifact.case_id, "artifact_exported", f"Artifact exported: {artifact.title} ({payload.format})")
        return ArtifactExportResult(artifact_id=artifact.id, format=payload.format, file_path=str(export_path))

    def execute_workflow(
        self,
        case_id: str,
        workflow_id: str,
        document_ids: list[str],
        engine: WorkflowEngine,
    ) -> dict:
        case = self.store.get("cases", case_id)
        if not case:
            raise KeyError(case_id)

        wf = engine.get_workflow(workflow_id)
        if not wf:
            raise KeyError(workflow_id)

        documents = self._get_documents(case_id, document_ids)

        workflow_run = WorkflowRunRecord(
            id=self.store.next_id("run"),
            case_id=case_id,
            workflow_type=wf.name,
            title=wf.name,
            status="running",
        )

        step_tasks = []
        task_ids = []
        for step in wf.steps:
            task = TaskRecord(
                id=self.store.next_id("task"),
                case_id=case_id,
                workflow_run_id=workflow_run.id,
                task_type=step.agent,
                title=step.label,
                document_ids=document_ids,
                status="pending",
            )
            self.store.create("tasks", task.model_dump())
            step_tasks.append(task)
            task_ids.append(task.id)

        workflow_run = workflow_run.model_copy(update={"task_ids": task_ids})
        self.store.create("workflow_runs", workflow_run.model_dump())
        self._append_timeline(case_id, "workflow_started", f"Workflow started: {wf.name}")

        # Execute first step immediately
        if step_tasks:
            first_task = step_tasks[0]
            executed, artifact = self.adapter.submit(case, first_task, documents)
            self.store.update("tasks", first_task.id, executed.model_dump())
            if artifact:
                self._store_artifact_once(artifact)
            step_tasks[0] = executed

        return {
            "workflow_run_id": workflow_run.id,
            "workflow_name": wf.name,
            "status": "running",
            "step_tasks": [t.model_dump() for t in step_tasks],
        }

    def _get_documents(self, case_id: str, document_ids: list[str]) -> list[DocumentRecord]:
        return [
            DocumentRecord(**doc)
            for doc in self.store.filter_by("documents", case_id=case_id)
            if doc["id"] in document_ids
        ]

    def _refresh_task(self, task: TaskRecord) -> TaskRecord:
        if task.status != "running":
            return task
        case = self.store.get("cases", task.case_id)
        if not case:
            return task
        documents = self._get_documents(task.case_id, task.document_ids)
        synced_task, artifact = self.adapter.sync(case, task, documents)
        self.store.update("tasks", synced_task.id, synced_task.model_dump())
        if artifact:
            self._store_artifact_once(artifact)
            self._append_timeline(task.case_id, "task_completed", f"Task completed: {task.title}")
        self._sync_workflow_run_from_task(synced_task)
        return synced_task

    def _sync_workflow_run_from_task(self, task: TaskRecord) -> None:
        workflow_run_record = self.store.get("workflow_runs", task.workflow_run_id)
        if not workflow_run_record:
            return
        workflow_run = WorkflowRunRecord(**workflow_run_record)
        status = "running"
        completed_at = ""
        if task.status == "completed":
            status = "completed"
            completed_at = utc_now()
        elif task.status == "failed":
            status = "failed"
        elif task.status == "canceled":
            status = "canceled"
        workflow_run = workflow_run.model_copy(
            update={
                "status": status,
                "error_message": task.error_message,
                "updated_at": utc_now(),
                "completed_at": completed_at,
            }
        )
        self.store.update("workflow_runs", workflow_run.id, workflow_run.model_dump())

    def _store_artifact_once(self, artifact: ArtifactRecord) -> None:
        if not self.store.exists("artifacts", artifact.id):
            self.store.create("artifacts", artifact.model_dump())

    def _append_timeline(self, case_id: str, kind: str, message: str) -> None:
        self.store.create(
            "timelines",
            TimelineItem(
                id=self.store.next_id("evt"),
                case_id=case_id,
                kind=kind,
                message=message,
            ).model_dump(),
        )
